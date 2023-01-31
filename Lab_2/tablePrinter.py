import numpy as np
import pandas as pd
import docx as docx

def genearateTable1(lat, lon, dist, time, vel, fileName):
    doc = docx.Document()
    valSize = np.arange(1, np.array(lat).size + 1)

    table_data = {'Index': valSize, 'Latitude': lat, 'Longitude': lon, 'Distance': dist, 'Time': time, 'Velocity': vel}
    df = pd.DataFrame(data=table_data)
    df['Time'] = pd.to_datetime(df.Time).dt.strftime('%H:%M:%S')

    # Initialise the table
    t = doc.add_table(rows=(df.shape[0] + 1), cols=df.shape[1])
    t.style = 'Light Shading Accent 1'

    # Add the column headings
    for j in range(df.shape[1]):
        t.cell(0, j).text = df.columns[j]

    # Add the body of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            cell = df.iat[i, j]
            t.cell(i + 1, j).text = str(cell)

    # Format Distance cells
    for i in range(df.shape[0]):
        cell = df.iat[i, 3]
        t.cell(i + 1, 3).text = str("%.2f" % cell)

    doc.save(fileName + '.docx')


def genearateTable2(totalDistance, averSpeed, maxFlight, minFlight, maxSpeed, minSpeed, timeFlight, fileName):
    doc = docx.Document()

    parameters = ['Total distance', 'Average velocity', 'Maximum altitude', 'Minimum altitude',
            'Maximum velocity', 'Minimum velocity', 'Total flight time']

    values = [totalDistance, velocityConverter(averSpeed),
            maxFlight, minFlight,
            velocityConverter(maxSpeed), velocityConverter(minSpeed),
            timeFlight]

    units = ['m', 'km/h', 'm', 'm', 'km/h', 'km/h', 's']

    index = np.arange(1, np.array(values).size + 1)

    table_data = {'Index': index, 'Parameter': parameters, 'Value': values, 'Units': units}

    df = pd.DataFrame(data=table_data)

    # Initialise the table
    t = doc.add_table(rows=(df.shape[0] + 1), cols=df.shape[1])
    t.style = 'Light Shading Accent 1'

    # Add the column headings
    for j in range(df.shape[1]):
        t.cell(0, j).text = df.columns[j]

    # Add the body of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            cell = df.iat[i, j]
            t.cell(i + 1, j).text = str(cell)

    cell = df.iat[1, 2]
    t.cell(2, 2).text = str("%.3f" % cell)

    doc.save(fileName + '.docx')

def velocityConverter(value):
    value = value / 1000 * 3600
    return value


def tableSeparator(v, type):
    match type:
        case "Table_1":
            genearateTable1(v._latitude, v._longtitude, v._arrDis, v._timeHigh, v._arrSpeed, type)
        case "Table_2":
            genearateTable2(v.getTotalDistance(), v.getAverSpeed(), v.getMaxAltitude(),
            v.getMinAltitude(), v._maxSpeed, v._minSpeed, v._timeFlight, type)
